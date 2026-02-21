"""API routes for Multi-Document Synthesis."""
from __future__ import annotations

import io
import json
from pathlib import Path
from typing import List, Optional
from pydantic import BaseModel, Field
from fastapi import APIRouter, Depends, Request, HTTPException, File, Form, UploadFile

from backend.app.services.security import require_api_key
from backend.app.services.synthesis.service import DocumentSynthesisService
from backend.app.schemas.synthesis import DocumentType, SynthesisRequest

router = APIRouter(dependencies=[Depends(require_api_key)])
MAX_SYNTHESIS_DOC_BYTES = 5 * 1024 * 1024

_EXTENSION_TO_DOC_TYPE = {
    ".txt": DocumentType.TEXT,
    ".md": DocumentType.TEXT,
    ".markdown": DocumentType.TEXT,
    ".csv": DocumentType.EXCEL,
    ".json": DocumentType.JSON,
    ".pdf": DocumentType.PDF,
    ".xlsx": DocumentType.EXCEL,
    ".xls": DocumentType.EXCEL,
    ".docx": DocumentType.WORD,
    ".doc": DocumentType.WORD,
}


class CreateSessionRequest(BaseModel):
    name: str = Field(..., min_length=1, max_length=200)


class AddDocumentRequest(BaseModel):
    name: str = Field(..., min_length=1, max_length=200)
    content: str = Field(..., min_length=10, max_length=5 * 1024 * 1024)
    doc_type: DocumentType = Field(default=DocumentType.TEXT)
    metadata: Optional[dict] = None


def _safe_decode_text(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore")


def _infer_document_type(filename: Optional[str], explicit: Optional[DocumentType]) -> DocumentType:
    if explicit:
        return explicit
    suffix = Path(filename or "").suffix.lower()
    return _EXTENSION_TO_DOC_TYPE.get(suffix, DocumentType.TEXT)


def _extract_pdf_text(file_bytes: bytes) -> str:
    try:
        import fitz
    except Exception as exc:
        raise HTTPException(status_code=500, detail="PDF extraction dependency not available") from exc

    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        chunks = []
        for page in doc:
            text = page.get_text("text") or ""
            if text.strip():
                chunks.append(text)
        doc.close()
        return "\n\n".join(chunks)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Could not parse PDF file") from exc


def _extract_excel_text(filename: str, file_bytes: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".csv":
        return _safe_decode_text(file_bytes)

    try:
        import pandas as pd
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail="Excel extraction dependencies are not available",
        ) from exc

    try:
        workbook = pd.ExcelFile(io.BytesIO(file_bytes))
        sheet_blocks = []
        for sheet_name in workbook.sheet_names[:20]:
            df = pd.read_excel(workbook, sheet_name=sheet_name)
            if df.empty:
                continue
            sheet_text = df.fillna("").to_csv(index=False, sep="\t").strip()
            if sheet_text:
                sheet_blocks.append(f"[Sheet: {sheet_name}]\n{sheet_text}")
        return "\n\n".join(sheet_blocks)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Could not parse Excel file") from exc


def _extract_word_text(filename: str, file_bytes: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix != ".docx":
        # .doc has no reliable pure-python parser in this stack; attempt plain decode fallback
        return _safe_decode_text(file_bytes)

    try:
        from docx import Document
    except Exception as exc:
        raise HTTPException(status_code=500, detail="DOCX extraction dependency not available") from exc

    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        table_lines = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    table_lines.append("\t".join(cells))
        parts = paragraphs + ([""] + table_lines if table_lines else [])
        return "\n".join(parts)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Could not parse Word file") from exc


def _extract_document_text(filename: str, file_bytes: bytes, doc_type: DocumentType) -> str:
    if doc_type == DocumentType.PDF:
        return _extract_pdf_text(file_bytes)
    if doc_type == DocumentType.EXCEL:
        return _extract_excel_text(filename, file_bytes)
    if doc_type == DocumentType.WORD:
        return _extract_word_text(filename, file_bytes)
    if doc_type == DocumentType.JSON:
        text = _safe_decode_text(file_bytes)
        try:
            parsed = json.loads(text)
            return json.dumps(parsed, indent=2, ensure_ascii=False)
        except Exception:
            return text
    return _safe_decode_text(file_bytes)


def get_service() -> DocumentSynthesisService:
    return DocumentSynthesisService()


@router.post("/sessions")
async def create_session(
    payload: CreateSessionRequest,
    request: Request,
    svc: DocumentSynthesisService = Depends(get_service),
):
    """Create a new synthesis session."""
    correlation_id = getattr(request.state, "correlation_id", None)
    session = svc.create_session(
        name=payload.name,
        correlation_id=correlation_id,
    )
    return {"status": "ok", "session": session.model_dump(mode="json"), "correlation_id": correlation_id}


@router.get("/sessions")
async def list_sessions(
    request: Request,
    svc: DocumentSynthesisService = Depends(get_service),
):
    """List all synthesis sessions."""
    correlation_id = getattr(request.state, "correlation_id", None)
    sessions = svc.list_sessions()
    return {
        "status": "ok",
        "sessions": [s.model_dump(mode="json") for s in sessions],
        "correlation_id": correlation_id,
    }


@router.get("/sessions/{session_id}")
async def get_session(
    session_id: str,
    request: Request,
    svc: DocumentSynthesisService = Depends(get_service),
):
    """Get a synthesis session by ID."""
    correlation_id = getattr(request.state, "correlation_id", None)
    session = svc.get_session(session_id)

    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    return {"status": "ok", "session": session.model_dump(mode="json"), "correlation_id": correlation_id}


@router.delete("/sessions/{session_id}")
async def delete_session(
    session_id: str,
    request: Request,
    svc: DocumentSynthesisService = Depends(get_service),
):
    """Delete a synthesis session."""
    correlation_id = getattr(request.state, "correlation_id", None)
    success = svc.delete_session(session_id)

    if not success:
        raise HTTPException(status_code=404, detail="Session not found")

    return {"status": "ok", "deleted": True, "correlation_id": correlation_id}


@router.post("/sessions/{session_id}/documents")
async def add_document(
    session_id: str,
    payload: AddDocumentRequest,
    request: Request,
    svc: DocumentSynthesisService = Depends(get_service),
):
    """Add a document to a synthesis session."""
    correlation_id = getattr(request.state, "correlation_id", None)
    document = svc.add_document(
        session_id=session_id,
        name=payload.name,
        content=payload.content,
        doc_type=payload.doc_type,
        metadata=payload.metadata,
        correlation_id=correlation_id,
    )

    if not document:
        raise HTTPException(status_code=404, detail="Session not found")

    return {"status": "ok", "document": document.model_dump(mode="json"), "correlation_id": correlation_id}


@router.post("/documents/extract")
async def extract_document(
    request: Request,
    file: UploadFile = File(...),
    doc_type: Optional[DocumentType] = Form(default=None),
):
    """Extract normalized text content from an uploaded file for synthesis."""
    correlation_id = getattr(request.state, "correlation_id", None)
    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")
    if len(file_bytes) > MAX_SYNTHESIS_DOC_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"File exceeds {MAX_SYNTHESIS_DOC_BYTES // (1024 * 1024)}MB limit",
        )

    filename = file.filename or "upload"
    inferred_type = _infer_document_type(filename, doc_type)
    extracted = _extract_document_text(filename, file_bytes, inferred_type).strip()
    if not extracted:
        raise HTTPException(status_code=400, detail="Could not extract text from uploaded file")

    truncated = False
    if len(extracted) > MAX_SYNTHESIS_DOC_BYTES:
        extracted = extracted[:MAX_SYNTHESIS_DOC_BYTES]
        truncated = True

    return {
        "status": "ok",
        "document": {
            "name": filename,
            "doc_type": inferred_type.value,
            "content": extracted,
            "truncated": truncated,
        },
        "correlation_id": correlation_id,
    }


@router.delete("/sessions/{session_id}/documents/{document_id}")
async def remove_document(
    session_id: str,
    document_id: str,
    request: Request,
    svc: DocumentSynthesisService = Depends(get_service),
):
    """Remove a document from a session."""
    correlation_id = getattr(request.state, "correlation_id", None)
    success = svc.remove_document(session_id, document_id)

    if not success:
        raise HTTPException(status_code=404, detail="Session or document not found")

    return {"status": "ok", "removed": True, "correlation_id": correlation_id}


@router.get("/sessions/{session_id}/inconsistencies")
async def find_inconsistencies(
    session_id: str,
    request: Request,
    svc: DocumentSynthesisService = Depends(get_service),
):
    """Find inconsistencies between documents in a session."""
    correlation_id = getattr(request.state, "correlation_id", None)

    session = svc.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    inconsistencies = svc.find_inconsistencies(session_id, correlation_id)

    return {
        "status": "ok",
        "inconsistencies": [i.model_dump(mode="json") for i in inconsistencies],
        "count": len(inconsistencies),
        "correlation_id": correlation_id,
    }


@router.post("/sessions/{session_id}/synthesize")
async def synthesize_documents(
    session_id: str,
    payload: SynthesisRequest,
    request: Request,
    svc: DocumentSynthesisService = Depends(get_service),
):
    """Synthesize information from all documents in a session."""
    correlation_id = getattr(request.state, "correlation_id", None)

    session = svc.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    if not session.documents:
        raise HTTPException(status_code=400, detail="No documents in session")

    result = svc.synthesize(session_id, payload, correlation_id)

    if not result:
        raise HTTPException(status_code=500, detail="Synthesis failed")

    return {"status": "ok", "result": result.model_dump(mode="json"), "correlation_id": correlation_id}
