"""DOCX rendering adapter using python-docx and html2docx."""

from __future__ import annotations

import logging
import time
from pathlib import Path
from typing import Optional

from backend.engine.domain.reports import OutputFormat
from .base import BaseRenderer, RenderContext, RenderResult

logger = logging.getLogger("neura.adapters.rendering.docx")


class DOCXRenderer(BaseRenderer):
    """Renderer that produces DOCX output from HTML."""

    def __init__(
        self,
        *,
        default_font_scale: float = 1.0,
    ) -> None:
        self._default_font_scale = default_font_scale

    @property
    def output_format(self) -> OutputFormat:
        return OutputFormat.DOCX

    def render(self, context: RenderContext) -> RenderResult:
        """Render DOCX from HTML."""
        start = time.perf_counter()

        try:
            self._ensure_output_dir(context.output_path)

            # Try to import required libraries
            try:
                from docx import Document
                from docx.shared import Inches, Pt
                from docx.enum.section import WD_ORIENT
            except ImportError:
                raise ImportError(
                    "python-docx is required for DOCX rendering. "
                    "Install with: pip install python-docx"
                )

            # Create document
            doc = Document()

            # Set landscape if requested
            if context.landscape:
                for section in doc.sections:
                    section.orientation = WD_ORIENT.LANDSCAPE
                    # Swap width and height for landscape
                    new_width = section.page_height
                    new_height = section.page_width
                    section.page_width = new_width
                    section.page_height = new_height

            # Convert HTML to DOCX content
            self._html_to_docx(doc, context.template_html, context.font_scale)

            # Save document
            doc.save(str(context.output_path))

            render_time = (time.perf_counter() - start) * 1000
            return RenderResult(
                success=True,
                output_path=context.output_path,
                format=OutputFormat.DOCX,
                size_bytes=self._get_file_size(context.output_path),
                render_time_ms=render_time,
            )
        except Exception as e:
            logger.exception("docx_render_failed")
            return RenderResult(
                success=False,
                output_path=None,
                format=OutputFormat.DOCX,
                error=str(e),
                render_time_ms=(time.perf_counter() - start) * 1000,
            )

    def _html_to_docx(
        self,
        doc,
        html: str,
        font_scale: Optional[float],
    ) -> None:
        """Convert HTML content to DOCX paragraphs.

        This is a simplified implementation. For production use,
        consider using a library like html2docx or mammoth.
        """
        try:
            from html2docx import html2docx
            html2docx(html, doc)
            return
        except ImportError:
            pass

        # Fallback: basic HTML parsing
        import re
        from html.parser import HTMLParser

        class SimpleHTMLParser(HTMLParser):
            def __init__(self, document):
                super().__init__()
                self.doc = document
                self.current_para = None
                self.in_table = False
                self.table_data = []
                self.current_row = []

            def handle_starttag(self, tag, attrs):
                if tag in ("p", "div"):
                    self.current_para = self.doc.add_paragraph()
                elif tag == "table":
                    self.in_table = True
                    self.table_data = []
                elif tag == "tr":
                    self.current_row = []
                elif tag == "br":
                    if self.current_para:
                        self.current_para.add_run("\n")
                elif tag in ("h1", "h2", "h3"):
                    level = int(tag[1])
                    self.current_para = self.doc.add_heading(level=level)

            def handle_endtag(self, tag):
                if tag == "table" and self.table_data:
                    self._create_table()
                    self.in_table = False
                elif tag == "tr" and self.current_row:
                    self.table_data.append(self.current_row)

            def handle_data(self, data):
                data = data.strip()
                if not data:
                    return
                if self.in_table:
                    self.current_row.append(data)
                elif self.current_para:
                    self.current_para.add_run(data)
                else:
                    self.current_para = self.doc.add_paragraph(data)

            def _create_table(self):
                if not self.table_data:
                    return
                rows = len(self.table_data)
                cols = max(len(row) for row in self.table_data)
                table = self.doc.add_table(rows=rows, cols=cols)
                table.style = "Table Grid"
                for i, row_data in enumerate(self.table_data):
                    for j, cell_data in enumerate(row_data):
                        if j < len(table.rows[i].cells):
                            table.rows[i].cells[j].text = str(cell_data)

        parser = SimpleHTMLParser(doc)
        parser.feed(html)


def render_docx_from_html(
    html_path: Path,
    output_path: Path,
    *,
    landscape: bool = False,
    font_scale: Optional[float] = None,
) -> RenderResult:
    """Convenience function to render DOCX from HTML file."""
    html_content = html_path.read_text(encoding="utf-8")

    renderer = DOCXRenderer()
    context = RenderContext(
        template_html=html_content,
        data={},
        output_format=OutputFormat.DOCX,
        output_path=output_path,
        landscape=landscape,
        font_scale=font_scale,
    )

    return renderer.render(context)
