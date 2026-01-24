"""
Document Ingestion Service
Handles file import, type detection, and document creation from various sources.
"""
from __future__ import annotations

import logging
import mimetypes
import hashlib
import zipfile
import io
import json
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, BinaryIO
from datetime import datetime, timezone
from enum import Enum

from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)


class FileType(str, Enum):
    """Supported file types for ingestion."""
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    XLSX = "xlsx"
    XLS = "xls"
    CSV = "csv"
    PPTX = "pptx"
    PPT = "ppt"
    TXT = "txt"
    RTF = "rtf"
    HTML = "html"
    MARKDOWN = "markdown"
    JSON = "json"
    XML = "xml"
    YAML = "yaml"
    IMAGE = "image"
    AUDIO = "audio"
    VIDEO = "video"
    ARCHIVE = "archive"
    UNKNOWN = "unknown"


class IngestionResult(BaseModel):
    """Result of document ingestion."""
    document_id: str
    filename: str
    file_type: FileType
    size_bytes: int
    pages: Optional[int] = None
    metadata: Dict[str, Any] = Field(default_factory=dict)
    preview_url: Optional[str] = None
    created_at: datetime = Field(default_factory=datetime.utcnow)
    processing_status: str = "completed"
    warnings: List[str] = Field(default_factory=list)


class BulkIngestionResult(BaseModel):
    """Result of bulk document ingestion."""
    total_files: int
    successful: int
    failed: int
    results: List[IngestionResult] = Field(default_factory=list)
    errors: List[Dict[str, str]] = Field(default_factory=list)


class StructuredDataImport(BaseModel):
    """Result of structured data import."""
    document_id: str
    table_name: str
    row_count: int
    column_count: int
    columns: List[str]
    sample_data: List[Dict[str, Any]]


class IngestionService:
    """
    Service for ingesting documents from various sources.
    Handles auto-detection, preview generation, and document creation.
    """

    # File extension to type mapping
    EXTENSION_MAP = {
        ".pdf": FileType.PDF,
        ".docx": FileType.DOCX,
        ".doc": FileType.DOC,
        ".xlsx": FileType.XLSX,
        ".xls": FileType.XLS,
        ".csv": FileType.CSV,
        ".pptx": FileType.PPTX,
        ".ppt": FileType.PPT,
        ".txt": FileType.TXT,
        ".rtf": FileType.RTF,
        ".html": FileType.HTML,
        ".htm": FileType.HTML,
        ".md": FileType.MARKDOWN,
        ".markdown": FileType.MARKDOWN,
        ".json": FileType.JSON,
        ".xml": FileType.XML,
        ".yaml": FileType.YAML,
        ".yml": FileType.YAML,
        ".png": FileType.IMAGE,
        ".jpg": FileType.IMAGE,
        ".jpeg": FileType.IMAGE,
        ".gif": FileType.IMAGE,
        ".bmp": FileType.IMAGE,
        ".webp": FileType.IMAGE,
        ".svg": FileType.IMAGE,
        ".mp3": FileType.AUDIO,
        ".wav": FileType.AUDIO,
        ".m4a": FileType.AUDIO,
        ".ogg": FileType.AUDIO,
        ".mp4": FileType.VIDEO,
        ".avi": FileType.VIDEO,
        ".mov": FileType.VIDEO,
        ".mkv": FileType.VIDEO,
        ".webm": FileType.VIDEO,
        ".zip": FileType.ARCHIVE,
        ".tar": FileType.ARCHIVE,
        ".gz": FileType.ARCHIVE,
        ".rar": FileType.ARCHIVE,
        ".7z": FileType.ARCHIVE,
    }

    def __init__(self):
        self._upload_dir: Optional[Path] = None

    def _get_upload_dir(self) -> Path:
        """Lazy load upload directory from settings."""
        if self._upload_dir is None:
            from backend.app.services.config import get_settings
            self._upload_dir = get_settings().uploads_dir
        return self._upload_dir

    def detect_file_type(self, filename: str, content: Optional[bytes] = None) -> FileType:
        """
        Detect file type from filename and optionally content.

        Args:
            filename: Original filename
            content: Optional file content for magic number detection

        Returns:
            Detected FileType
        """
        ext = Path(filename).suffix.lower()

        if ext in self.EXTENSION_MAP:
            return self.EXTENSION_MAP[ext]

        # Try MIME type detection
        mime_type, _ = mimetypes.guess_type(filename)
        if mime_type:
            if mime_type.startswith("image/"):
                return FileType.IMAGE
            elif mime_type.startswith("audio/"):
                return FileType.AUDIO
            elif mime_type.startswith("video/"):
                return FileType.VIDEO
            elif mime_type.startswith("text/"):
                return FileType.TXT

        # Magic number detection for content
        if content:
            if content[:4] == b"%PDF":
                return FileType.PDF
            elif content[:4] == b"PK\x03\x04":
                # Could be DOCX, XLSX, PPTX, or ZIP
                return self._detect_office_type(content)
            elif content[:3] == b"\xef\xbb\xbf" or content[:1000].decode("utf-8", errors="ignore").strip().startswith(("{", "[")):
                return FileType.JSON

        return FileType.UNKNOWN

    def _detect_office_type(self, content: bytes) -> FileType:
        """Detect specific Office format from ZIP-based file."""
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                names = zf.namelist()
                if any("word/" in n for n in names):
                    return FileType.DOCX
                elif any("xl/" in n for n in names):
                    return FileType.XLSX
                elif any("ppt/" in n for n in names):
                    return FileType.PPTX
        except zipfile.BadZipFile:
            pass
        return FileType.ARCHIVE

    async def ingest_file(
        self,
        filename: str,
        content: bytes,
        metadata: Optional[Dict[str, Any]] = None,
        auto_ocr: bool = True,
        generate_preview: bool = True,
    ) -> IngestionResult:
        """
        Ingest a single file and create a document.

        Args:
            filename: Original filename
            content: File content
            metadata: Optional metadata to attach
            auto_ocr: Whether to OCR scanned documents
            generate_preview: Whether to generate preview images

        Returns:
            IngestionResult with document details
        """
        file_type = self.detect_file_type(filename, content)
        doc_id = self._generate_document_id(filename, content)

        # Save file
        upload_dir = self._get_upload_dir()
        file_path = upload_dir / doc_id / filename
        file_path.parent.mkdir(parents=True, exist_ok=True)
        file_path.write_bytes(content)

        # Extract metadata based on file type
        extracted_metadata = await self._extract_metadata(file_path, file_type)
        if metadata:
            extracted_metadata.update(metadata)

        # Get page count for documents
        pages = await self._get_page_count(file_path, file_type)

        # Generate preview if requested
        preview_url = None
        if generate_preview:
            preview_url = await self._generate_preview(file_path, file_type, doc_id)

        # OCR if needed
        warnings = []
        if auto_ocr and file_type == FileType.PDF:
            is_scanned = await self._is_scanned_pdf(file_path)
            if is_scanned:
                await self._perform_ocr(file_path)
                warnings.append("Document was OCR'd from scanned images")

        return IngestionResult(
            document_id=doc_id,
            filename=filename,
            file_type=file_type,
            size_bytes=len(content),
            pages=pages,
            metadata=extracted_metadata,
            preview_url=preview_url,
            warnings=warnings,
        )

    async def ingest_from_url(
        self,
        url: str,
        filename: Optional[str] = None,
    ) -> IngestionResult:
        """
        Ingest a document from a URL.

        Args:
            url: URL to download from
            filename: Optional filename override

        Returns:
            IngestionResult with document details
        """
        import aiohttp

        async with aiohttp.ClientSession() as session:
            async with session.get(url) as response:
                response.raise_for_status()
                content = await response.read()

                # Determine filename
                if not filename:
                    # Try to get from Content-Disposition header
                    cd = response.headers.get("Content-Disposition", "")
                    if "filename=" in cd:
                        filename = cd.split("filename=")[1].strip('"')
                    else:
                        # Use URL path
                        from urllib.parse import urlparse
                        filename = Path(urlparse(url).path).name or "downloaded_file"

                return await self.ingest_file(filename, content)

    async def ingest_zip_archive(
        self,
        filename: str,
        content: bytes,
        preserve_structure: bool = True,
        flatten: bool = False,
    ) -> BulkIngestionResult:
        """
        Ingest documents from a ZIP archive.

        Args:
            filename: Archive filename
            content: ZIP file content
            preserve_structure: Preserve folder structure
            flatten: Flatten all files to root

        Returns:
            BulkIngestionResult with all ingested documents
        """
        results = []
        errors = []

        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                for name in zf.namelist():
                    # Skip directories
                    if name.endswith("/"):
                        continue

                    # Skip hidden/system files
                    if any(part.startswith(".") or part.startswith("__") for part in Path(name).parts):
                        continue

                    try:
                        file_content = zf.read(name)
                        file_name = Path(name).name if flatten else name

                        result = await self.ingest_file(
                            filename=file_name,
                            content=file_content,
                            metadata={"source_archive": filename, "original_path": name},
                        )
                        results.append(result)
                    except Exception as e:
                        errors.append({"file": name, "error": str(e)})
                        logger.error(f"Failed to ingest {name} from archive: {e}")

        except zipfile.BadZipFile as e:
            errors.append({"file": filename, "error": f"Invalid ZIP file: {e}"})

        return BulkIngestionResult(
            total_files=len(results) + len(errors),
            successful=len(results),
            failed=len(errors),
            results=results,
            errors=errors,
        )

    async def import_structured_data(
        self,
        filename: str,
        content: bytes,
        format_hint: Optional[str] = None,
    ) -> StructuredDataImport:
        """
        Import structured data (JSON, XML, YAML) as an editable table.

        Args:
            filename: Original filename
            content: File content
            format_hint: Optional format hint

        Returns:
            StructuredDataImport with table details
        """
        file_type = format_hint or self.detect_file_type(filename, content)
        text_content = content.decode("utf-8")

        data = None
        if file_type in (FileType.JSON, "json"):
            data = json.loads(text_content)
        elif file_type in (FileType.YAML, "yaml"):
            import yaml
            data = yaml.safe_load(text_content)
        elif file_type in (FileType.XML, "xml"):
            import xml.etree.ElementTree as ET
            root = ET.fromstring(text_content)
            data = self._xml_to_dict(root)

        # Normalize to list of dicts
        if isinstance(data, dict):
            # Check if it's a single record or contains a list
            for key, value in data.items():
                if isinstance(value, list) and all(isinstance(v, dict) for v in value):
                    data = value
                    break
            else:
                data = [data]
        elif not isinstance(data, list):
            data = [{"value": data}]

        # Get columns
        columns = []
        for record in data[:10]:
            if isinstance(record, dict):
                for key in record.keys():
                    if key not in columns:
                        columns.append(key)

        # Create document
        doc_id = self._generate_document_id(filename, content)

        return StructuredDataImport(
            document_id=doc_id,
            table_name=Path(filename).stem,
            row_count=len(data),
            column_count=len(columns),
            columns=columns,
            sample_data=data[:10],
        )

    def _xml_to_dict(self, element) -> Dict[str, Any]:
        """Convert XML element to dictionary."""
        result = {}
        for child in element:
            if len(child) == 0:
                result[child.tag] = child.text
            else:
                result[child.tag] = self._xml_to_dict(child)
        if element.attrib:
            result["@attributes"] = element.attrib
        return result

    def _generate_document_id(self, filename: str, content: bytes) -> str:
        """Generate unique document ID."""
        hash_input = f"{filename}:{len(content)}:{datetime.now(timezone.utc).isoformat()}"
        return hashlib.sha256(hash_input.encode()).hexdigest()[:16]

    async def _extract_metadata(self, file_path: Path, file_type: FileType) -> Dict[str, Any]:
        """Extract metadata from file."""
        metadata = {
            "original_filename": file_path.name,
            "file_type": file_type.value,
            "ingested_at": datetime.now(timezone.utc).isoformat(),
        }

        try:
            if file_type == FileType.PDF:
                import fitz
                doc = fitz.open(str(file_path))
                pdf_metadata = doc.metadata
                if pdf_metadata:
                    metadata.update({
                        "title": pdf_metadata.get("title"),
                        "author": pdf_metadata.get("author"),
                        "subject": pdf_metadata.get("subject"),
                        "creator": pdf_metadata.get("creator"),
                        "creation_date": pdf_metadata.get("creationDate"),
                    })
                doc.close()

            elif file_type == FileType.DOCX:
                from docx import Document
                doc = Document(str(file_path))
                props = doc.core_properties
                metadata.update({
                    "title": props.title,
                    "author": props.author,
                    "subject": props.subject,
                    "created": props.created.isoformat() if props.created else None,
                    "modified": props.modified.isoformat() if props.modified else None,
                })

            elif file_type == FileType.IMAGE:
                from PIL import Image
                with Image.open(file_path) as img:
                    metadata.update({
                        "width": img.width,
                        "height": img.height,
                        "format": img.format,
                        "mode": img.mode,
                    })

        except Exception as e:
            logger.warning(f"Failed to extract metadata: {e}")

        return {k: v for k, v in metadata.items() if v is not None}

    async def _get_page_count(self, file_path: Path, file_type: FileType) -> Optional[int]:
        """Get page count for documents."""
        try:
            if file_type == FileType.PDF:
                import fitz
                doc = fitz.open(str(file_path))
                count = len(doc)
                doc.close()
                return count

            elif file_type == FileType.DOCX:
                # Approximate from paragraphs (rough estimate)
                from docx import Document
                doc = Document(str(file_path))
                # ~40 paragraphs per page is a rough estimate
                return max(1, len(doc.paragraphs) // 40)

            elif file_type == FileType.PPTX:
                from pptx import Presentation
                prs = Presentation(str(file_path))
                return len(prs.slides)

        except Exception as e:
            logger.warning(f"Failed to get page count: {e}")

        return None

    async def _generate_preview(self, file_path: Path, file_type: FileType, doc_id: str) -> Optional[str]:
        """Generate preview image for document."""
        try:
            preview_dir = self._get_upload_dir() / doc_id / "previews"
            preview_dir.mkdir(parents=True, exist_ok=True)
            preview_path = preview_dir / "preview.png"

            if file_type == FileType.PDF:
                import fitz
                doc = fitz.open(str(file_path))
                if len(doc) > 0:
                    page = doc[0]
                    pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
                    pix.save(str(preview_path))
                doc.close()

            elif file_type == FileType.IMAGE:
                from PIL import Image
                with Image.open(file_path) as img:
                    # Create thumbnail
                    img.thumbnail((400, 400))
                    img.save(preview_path, "PNG")

            if preview_path.exists():
                return f"/uploads/{doc_id}/previews/preview.png"

        except Exception as e:
            logger.warning(f"Failed to generate preview: {e}")

        return None

    async def _is_scanned_pdf(self, file_path: Path) -> bool:
        """Check if PDF appears to be scanned (image-based)."""
        try:
            import fitz
            doc = fitz.open(str(file_path))
            if len(doc) == 0:
                doc.close()
                return False

            # Check first few pages
            for page_num in range(min(3, len(doc))):
                page = doc[page_num]
                text = page.get_text()
                if len(text.strip()) > 50:
                    doc.close()
                    return False

            doc.close()
            return True
        except Exception:
            return False

    async def _perform_ocr(self, file_path: Path) -> None:
        """Perform OCR on a scanned PDF."""
        # This would integrate with the DocAI OCR service
        logger.info(f"OCR would be performed on: {file_path}")


# Singleton instance
ingestion_service = IngestionService()
