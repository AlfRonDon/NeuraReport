from __future__ import annotations

import contextlib
import logging
import os
import re
from io import BytesIO
from pathlib import Path
from typing import Iterable, Optional

try:
    from html2docx import html2docx  # type: ignore
except ImportError:  # pragma: no cover
    html2docx = None  # type: ignore

try:  # pragma: no cover - only exercised when python-docx is available
    from docx import Document  # type: ignore
    from docx.enum.section import WD_ORIENT  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.shared import Mm, Pt  # type: ignore
except ImportError:  # pragma: no cover
    Document = None  # type: ignore
    WD_ORIENT = None  # type: ignore
    WD_ALIGN_PARAGRAPH = None  # type: ignore
    Mm = None  # type: ignore
    Pt = None  # type: ignore

try:  # pragma: no cover - optional dependency for PDF conversion
    from pdf2docx import Converter  # type: ignore
except ImportError:  # pragma: no cover
    Converter = None  # type: ignore

try:
    from lxml import etree
    from lxml import html as lxml_html  # type: ignore
except ImportError:  # pragma: no cover
    etree = None  # type: ignore
    lxml_html = None  # type: ignore

from .html_table_parser import extract_tables

logger = logging.getLogger("neura.reports.docx")

_BODY_TAG_RE = re.compile(r"(?is)<body\b(?P<attrs>[^>]*)>", re.MULTILINE)
_STYLE_ATTR_RE = re.compile(r'(?is)(style\s*=\s*)(["\'])(?P<value>.*?)\2')
_STYLE_BLOCK_RE = re.compile(r"(?is)<style\b[^>]*>.*?</style>")
_SCRIPT_BLOCK_RE = re.compile(r"(?is)<script\b[^>]*>.*?</script>")
_TITLE_RE = re.compile(r"(?is)<title\b[^>]*>(?P<value>.*?)</title>")
_TAG_RE = re.compile(r"(?is)<[^>]+>")
_WHITESPACE_RE = re.compile(r"\s+")

_A4_LANDSCAPE_WIDTH_MM = 297
_A4_LANDSCAPE_HEIGHT_MM = 210
_MAX_FALLBACK_ROWS = 500
_MAX_FALLBACK_TABLES = 8


def _inject_body_style(html_text: str, style_rule: str) -> str:
    """Attach/extend a style attribute on the <body> tag without disturbing the markup."""
    match = _BODY_TAG_RE.search(html_text)
    if not match:
        return f'<body style="{style_rule}">{html_text}</body>'

    attrs = match.group("attrs") or ""
    new_attrs = attrs

    def _style_repl(style_match: re.Match[str]) -> str:
        existing = (style_match.group("value") or "").strip().rstrip(";")
        merged = "; ".join(filter(None, [existing, style_rule]))
        return f"{style_match.group(1)}{style_match.group(2)}{merged}{style_match.group(2)}"

    if _STYLE_ATTR_RE.search(attrs):
        new_attrs = _STYLE_ATTR_RE.sub(_style_repl, attrs, count=1)
    else:
        spacer = "" if attrs.endswith(" ") or not attrs else " "
        new_attrs = f'{attrs}{spacer}style="{style_rule}"'

    new_tag = f"<body{new_attrs}>"
    return f"{html_text[: match.start()]}{new_tag}{html_text[match.end():]}"


def _apply_body_font_scale(html_text: str, scale: float | None) -> str:
    if not scale or scale <= 0:
        return html_text

    clamped = max(0.5, min(scale, 1.0))
    percent = round(clamped * 100, 1)
    style_rule = f"font-size: {percent}%; line-height: 1.15;"
    return _inject_body_style(html_text, style_rule)


def _append_inline_style(node, style: str) -> None:
    if node is None or not style:
        return
    existing = (node.get("style") or "").strip()
    if existing and not existing.endswith(";"):
        existing = f"{existing};"
    parts = [part.strip() for part in (existing.rstrip(";"), style.strip()) if part and part.strip()]
    node.set("style", "; ".join(parts))


def _inline_report_styles(html_text: str) -> str:
    if not html_text:
        return html_text
    try:
        document = lxml_html.fromstring(html_text)
    except Exception:
        return html_text

    def _set_style(xpath: str, style: str) -> None:
        if not style:
            return
        for node in document.xpath(xpath):
            _append_inline_style(node, style)

    _set_style(
        "//body", "font-family: 'Times New Roman', serif; font-size: 12px; color: #000; line-height: 1.2; margin: 0;"
    )
    _set_style("//div[@id='report-header']", "margin-top: 0; page-break-inside: avoid;")
    _set_style(
        "//div[@id='report-header']//table",
        "width: 100%; border: 1px solid #000; border-collapse: collapse; table-layout: fixed;",
    )
    _set_style("//div[@id='report-header']//td", "vertical-align: top; padding: 1.6mm 2.4mm; border: 1px solid #000;")
    _set_style(
        "//div[contains(concat(' ', normalize-space(@class), ' '), ' title-wrap ')]",
        "margin: 3mm 0 2.5mm 0; text-align: center; font-weight: bold; font-size: 18px; border-top: 1px solid #000; border-bottom: 1px solid #000; padding: 2mm 0;",
    )
    _set_style("//table[@id='data-table']", "width: 100%; border-collapse: collapse; table-layout: fixed;")
    _set_style(
        "//table[@id='data-table']//th | //table[@id='data-table']//td", "border: 1px solid #000; padding: 1mm 2.2mm;"
    )
    _set_style(
        "//table[@id='data-table']//thead//th",
        "text-align: center; font-weight: bold; white-space: nowrap; padding: 1.6mm 2.4mm;",
    )
    _set_style(
        "//table[@id='data-table']//td[contains(concat(' ', normalize-space(@class), ' '), ' num ')]",
        "text-align: right;",
    )
    _set_style("//tfoot[@id='report-totals']//td", "font-weight: bold; border-top: 1.2px solid #000;")
    _set_style(
        "//tfoot[@id='report-totals']//td[contains(concat(' ', normalize-space(@class), ' '), ' label ')]",
        "text-align: left;",
    )
    _set_style(
        "//footer[@id='report-footer']",
        "font-size: 11px; color: #000; display: flex; justify-content: space-between; align-items: center;",
    )
    _set_style(
        "//footer[@id='report-footer']//div[contains(concat(' ', normalize-space(@class), ' '), ' page ')]",
        "text-align: center;",
    )

    return etree.tostring(document, encoding="unicode", method="html")


def _extract_report_title(html_text: str) -> str:
    if not html_text:
        return ""
    try:
        document = lxml_html.fromstring(html_text)
    except Exception:
        return _extract_html_title(html_text)
    title_nodes = document.xpath("//*[contains(concat(' ', normalize-space(@class), ' '), ' title-wrap ')]")
    if title_nodes:
        return _normalize_whitespace(title_nodes[0].text_content())
    return _extract_html_title(html_text)


def _extract_footer_brand(html_text: str) -> str:
    if not html_text:
        return ""
    try:
        document = lxml_html.fromstring(html_text)
    except Exception:
        return ""
    brand_nodes = document.xpath(
        "//footer[@id='report-footer']//div[contains(concat(' ', normalize-space(@class), ' '), ' brand ')]"
    )
    if brand_nodes:
        return _normalize_whitespace(brand_nodes[0].text_content())
    return ""


def _configure_document_layout(document, *, body_font_scale: float | None = None) -> None:
    if document is None:
        return
    if Mm is not None:
        try:
            section = document.sections[0]
        except Exception:
            section = None
        if section is not None:
            section.left_margin = Mm(16)
            section.right_margin = Mm(16)
            section.top_margin = Mm(14)
            section.bottom_margin = Mm(14)


def _infer_numeric_columns(header_row: list[str]) -> set[int]:
    numeric_columns: set[int] = set()
    tokens = ("wt", "weight", "error", "%", "kg", "total", "qty")
    for idx, cell in enumerate(header_row or []):
        text = (cell or "").lower()
        if idx == 0:
            continue
        if any(token in text for token in tokens):
            numeric_columns.add(idx)
    return numeric_columns


def _column_widths(max_columns: int, *, ratios: Optional[Iterable[float]], document) -> list[float] | None:
    if max_columns <= 0 or document is None or Mm is None:
        return None
    try:
        section = document.sections[0]
        available = section.page_width - section.left_margin - section.right_margin
    except Exception:
        return None
    if available <= 0:
        return None
    ratio_list = list(ratios or [])
    if ratio_list and len(ratio_list) < max_columns:
        last = ratio_list[-1]
        ratio_list.extend([last] * (max_columns - len(ratio_list)))
    if not ratio_list:
        ratio_list = [1.0] * max_columns
    total = sum(ratio_list) or 1.0
    return [available * (value / total) for value in ratio_list[:max_columns]]


def _write_docx_table(
    document,
    rows: list[list[str]],
    *,
    header_rows: int = 0,
    column_widths: Optional[Iterable[float]] = None,
    numeric_columns: Optional[set[int]] = None,
) -> None:
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = document.add_table(rows=len(rows), cols=max_cols)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    table.autofit = False
    widths = _column_widths(max_cols, ratios=column_widths, document=document)

    for r_idx, row in enumerate(rows):
        for c_idx in range(max_cols):
            value = row[c_idx] if c_idx < len(row) else ""
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.text = _normalize_whitespace(value)
            if WD_ALIGN_PARAGRAPH is not None:
                if numeric_columns and c_idx in (numeric_columns or set()) and r_idx >= header_rows:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif r_idx < header_rows:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = r_idx < header_rows

    if widths:
        for c_idx, width in enumerate(widths):
            for cell in table.columns[c_idx].cells:
                cell.width = width


def _strip_style_blocks(html_text: str) -> str:
    cleaned = _STYLE_BLOCK_RE.sub("", html_text or "")
    cleaned = _SCRIPT_BLOCK_RE.sub("", cleaned)
    return cleaned


def _normalize_whitespace(text: str) -> str:
    return _WHITESPACE_RE.sub(" ", text or "").strip()


def _strip_html_to_text(html_text: str) -> str:
    return _normalize_whitespace(_TAG_RE.sub(" ", html_text or ""))


def _extract_html_title(html_text: str) -> str:
    match = _TITLE_RE.search(html_text or "")
    if not match:
        return ""
    return _strip_html_to_text(match.group("value"))


def _extract_section_nodes(html_text: str) -> list[etree._Element]:
    try:
        document = lxml_html.fromstring(html_text or "")
    except Exception:
        return []
    sections = document.xpath("//div[contains(concat(' ', normalize-space(@class), ' '), ' nr-key-section ')]")
    if sections:
        return sections
    body = document.xpath("//body")
    return body or [document]


def _fallback_docx_from_tables(html_text: str, output_path: Path, *, body_font_scale: float | None) -> Optional[Path]:
    if Document is None:  # pragma: no cover
        return None

    sections = _extract_section_nodes(html_text)
    if not sections:
        try:
            sections = [lxml_html.fromstring(html_text or "<div></div>")]
        except Exception:
            sections = []
    if not sections:
        return None

    try:
        document = Document()  # type: ignore[call-arg]
    except Exception as exc:  # pragma: no cover
        logger.warning(
            "docx_fallback_init_failed",
            extra={
                "event": "docx_fallback_init_failed",
                "error": str(exc),
            },
        )
        return None

    _configure_document_layout(document, body_font_scale=body_font_scale)

    for index, section_node in enumerate(sections):
        if index > 0:
            document.add_page_break()

        section_html = etree.tostring(section_node, encoding="unicode", method="html")
        title_text = _extract_report_title(section_html) or _extract_report_title(html_text)
        if title_text:
            paragraph = document.add_paragraph(title_text)
            if WD_ALIGN_PARAGRAPH is not None:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True

        tables = extract_tables(section_html, max_tables=_MAX_FALLBACK_TABLES)
        if not tables:
            plain_text = _strip_html_to_text(section_html)
            document.add_paragraph(plain_text or "Report data unavailable.")
            continue

        header_rows = tables[0] if len(tables) >= 1 else []
        data_rows = tables[1] if len(tables) >= 2 else []
        extra_tables = tables[2:] if len(tables) > 2 else []

        if header_rows:
            capped_header = header_rows[:_MAX_FALLBACK_ROWS]
            _write_docx_table(
                document,
                capped_header,
                header_rows=0,
                column_widths=[2.5, 1.5],
            )
            document.add_paragraph("")

        if data_rows:
            capped_rows = data_rows[:_MAX_FALLBACK_ROWS]
            numeric_columns = _infer_numeric_columns(capped_rows[0] if capped_rows else [])
            _write_docx_table(
                document,
                capped_rows,
                header_rows=1,
                numeric_columns=numeric_columns,
            )
            document.add_paragraph("")

        for rows in extra_tables:
            if not rows:
                continue
            capped_rows = rows[:_MAX_FALLBACK_ROWS]
            _write_docx_table(document, capped_rows, header_rows=1)
            document.add_paragraph("")

        brand_text = _extract_footer_brand(section_html) or _extract_footer_brand(html_text)
        if brand_text:
            footer_paragraph = document.add_paragraph(brand_text)
            if WD_ALIGN_PARAGRAPH is not None:
                footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    try:
        document.save(output_path)
    except Exception as exc:  # pragma: no cover
        logger.warning(
            "docx_fallback_save_failed",
            extra={
                "event": "docx_fallback_save_failed",
                "docx_path": str(output_path),
                "error": str(exc),
            },
        )
        return None
    return output_path


def _clamp_body_scale(scale: float | None) -> float:
    if not scale or not isinstance(scale, (int, float)):
        return 1.0
    return max(0.3, min(float(scale), 1.0))


def _enforce_landscape_layout(docx_path: Path, *, margin_mm: float = 10.0) -> None:
    if Document is None or WD_ORIENT is None or Mm is None:  # pragma: no cover
        logger.debug(
            "docx_landscape_skipped",
            extra={
                "event": "docx_landscape_skipped",
                "reason": "python-docx unavailable",
                "docx_path": str(docx_path),
            },
        )
        return

    try:
        document = Document(docx_path)  # type: ignore[call-arg]
    except Exception as exc:  # pragma: no cover
        logger.warning(
            "docx_landscape_open_failed",
            extra={
                "event": "docx_landscape_open_failed",
                "docx_path": str(docx_path),
                "error": str(exc),
            },
        )
        return

    width = Mm(_A4_LANDSCAPE_WIDTH_MM)
    height = Mm(_A4_LANDSCAPE_HEIGHT_MM)
    margin = Mm(margin_mm)

    for section in document.sections:
        section.orientation = WD_ORIENT.LANDSCAPE  # type: ignore[assignment]
        section.page_width = width
        section.page_height = height
        section.left_margin = margin
        section.right_margin = margin
        section.top_margin = margin
        section.bottom_margin = margin

    try:
        document.save(docx_path)
    except Exception as exc:  # pragma: no cover
        logger.warning(
            "docx_landscape_save_failed",
            extra={
                "event": "docx_landscape_save_failed",
                "docx_path": str(docx_path),
                "error": str(exc),
            },
        )


def html_file_to_docx(
    html_path: Path,
    output_path: Path,
    *,
    landscape: bool = False,
    body_font_scale: float | None = None,
) -> Optional[Path]:
    """
    Convert an HTML file into a DOCX document using html2docx.

    Parameters
    ----------
    html_path:
        Source HTML document to convert.
    output_path:
        Where the generated DOCX should be written.
    landscape:
        When True, enforces an A4 landscape layout inside the resulting DOCX.
    body_font_scale:
        Optional percentage (0-1] used to downscale the body font size before conversion.

    Returns the output path on success, or None when conversion is unavailable.
    """
    if lxml_html is None or etree is None:
        logger.warning("lxml not available — DOCX export disabled")
        return None

    html_text = html_path.read_text(encoding="utf-8", errors="ignore")
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    html_with_inline = _inline_report_styles(html_text)
    html_for_docx = _apply_body_font_scale(_strip_style_blocks(html_with_inline), body_font_scale)

    if html2docx is not None:  # pragma: no branch
        try:
            title_text = _extract_html_title(html_for_docx) or "Report"
            buffer: BytesIO = html2docx(html_for_docx, title_text)  # type: ignore[call-arg]
        except Exception as exc:  # pragma: no cover
            logger.exception(
                "docx_export_html2docx_failed",
                extra={
                    "event": "docx_export_html2docx_failed",
                    "html_path": str(html_path),
                    "docx_path": str(output_path),
                    "error": str(exc),
                },
            )
        else:
            with output_path.open("wb") as handle:
                handle.write(buffer.getvalue())
            if landscape:
                _enforce_landscape_layout(output_path)
            logger.info(
                "docx_export_success",
                extra={
                    "event": "docx_export_success",
                    "html_path": str(html_path),
                    "docx_path": str(output_path),
                    "landscape": landscape,
                    "font_scale": body_font_scale,
                    "strategy": "html2docx",
                },
            )
            return output_path

    structured = _fallback_docx_from_tables(html_text, output_path, body_font_scale=body_font_scale)
    if structured:
        if landscape:
            _enforce_landscape_layout(output_path)
        logger.info(
            "docx_export_success",
            extra={
                "event": "docx_export_success",
                "html_path": str(html_path),
                "docx_path": str(output_path),
                "landscape": landscape,
                "font_scale": body_font_scale,
                "strategy": "structured",
            },
        )
        return output_path

    logger.warning(
        "docx_export_unavailable",
        extra={
            "event": "docx_export_unavailable",
            "reason": "python-docx unavailable" if Document is None else "html2docx not installed",
            "html_path": str(html_path),
        },
    )
    return None


_PDF2DOCX_TIMEOUT = int(os.environ.get("NEURA_PDF2DOCX_TIMEOUT", "120"))  # seconds


def pdf_file_to_docx(
    pdf_path: Path,
    output_path: Path,
    *,
    start_page: int = 0,
    end_page: int | None = None,
) -> Optional[Path]:
    """
    Convert an already-rendered PDF into DOCX using pdf2docx for near-carbon-copy layout.
    Returns None when conversion is unavailable or fails so callers can fall back to HTML export.

    Uses a timeout (default 120s) because pdf2docx can hang indefinitely on
    complex Excel-based PDFs during its "[2/4] Analyzing document" phase.
    """
    if Converter is None:  # pragma: no cover
        logger.debug(
            "docx_pdf_convert_skipped",
            extra={
                "event": "docx_pdf_convert_skipped",
                "reason": "pdf2docx unavailable",
                "pdf_path": str(pdf_path),
            },
        )
        return None

    pdf_path = Path(pdf_path)
    if not pdf_path.exists():
        logger.warning(
            "docx_pdf_convert_missing_pdf",
            extra={
                "event": "docx_pdf_convert_missing_pdf",
                "pdf_path": str(pdf_path),
            },
        )
        return None

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        converter = Converter(str(pdf_path))
    except Exception as exc:  # pragma: no cover
        logger.warning(
            "docx_pdf_convert_open_failed",
            extra={
                "event": "docx_pdf_convert_open_failed",
                "pdf_path": str(pdf_path),
                "error": str(exc),
            },
        )
        return None

    # Run the conversion in a separate thread with a timeout to prevent
    # indefinite hangs on complex PDFs (pdf2docx can stall for 15+ min).
    import concurrent.futures

    def _do_convert():
        converter.convert(str(output_path), start=start_page, end=end_page)

    try:
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
            fut = pool.submit(_do_convert)
            fut.result(timeout=_PDF2DOCX_TIMEOUT)
    except concurrent.futures.TimeoutError:
        logger.warning(
            "docx_pdf_convert_timeout",
            extra={
                "event": "docx_pdf_convert_timeout",
                "pdf_path": str(pdf_path),
                "timeout_seconds": _PDF2DOCX_TIMEOUT,
            },
        )
        with contextlib.suppress(Exception):
            converter.close()
        with contextlib.suppress(FileNotFoundError):
            output_path.unlink(missing_ok=True)
        return None
    except Exception as exc:  # pragma: no cover
        logger.warning(
            "docx_pdf_convert_failed",
            extra={
                "event": "docx_pdf_convert_failed",
                "pdf_path": str(pdf_path),
                "docx_path": str(output_path),
                "error": str(exc),
            },
        )
        return None
    finally:
        with contextlib.suppress(Exception):
            converter.close()

    logger.info(
        "docx_pdf_convert_success",
        extra={
            "event": "docx_pdf_convert_success",
            "pdf_path": str(pdf_path),
            "docx_path": str(output_path),
            "start_page": start_page,
            "end_page": end_page,
        },
    )
    return output_path
