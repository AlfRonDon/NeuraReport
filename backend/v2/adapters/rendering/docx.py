"""
DOCX renderer from HTML.
"""

from __future__ import annotations

import logging
from pathlib import Path

from .base import Renderer
from ...core import Result, Ok, Err, DomainError, RenderError

logger = logging.getLogger("neura.adapters.rendering.docx")


class HTMLToDocxRenderer(Renderer):
    """
    Convert HTML to DOCX using python-docx and htmldocx.
    """

    async def render(
        self,
        source: Path,
        output: Path,
        **options,
    ) -> Result[Path, DomainError]:
        """Render HTML file to DOCX."""
        if not source.exists():
            return Err(RenderError(
                code="source_not_found",
                message=f"Source file not found: {source}",
                format="docx",
                stage="load",
            ))

        try:
            content = source.read_text(encoding="utf-8")
            return await self.render_from_string(content, output, **options)
        except Exception as e:
            return Err(RenderError(
                code="read_failed",
                message=f"Failed to read source file: {e}",
                format="docx",
                stage="load",
                cause=e,
            ))

    async def render_from_string(
        self,
        content: str,
        output: Path,
        **options,
    ) -> Result[Path, DomainError]:
        """Render HTML string to DOCX."""
        try:
            from docx import Document
        except ImportError:
            return Err(RenderError(
                code="docx_not_installed",
                message="python-docx is not installed. Run: pip install python-docx",
                format="docx",
                stage="init",
            ))

        try:
            from htmldocx import HtmlToDocx
        except ImportError:
            # Fall back to basic conversion
            return await self._render_basic(content, output, options)

        output.parent.mkdir(parents=True, exist_ok=True)

        try:
            doc = Document()
            converter = HtmlToDocx()

            # Add HTML content
            converter.add_html_to_document(content, doc)

            # Save the document
            doc.save(str(output))

            if not output.exists():
                return Err(RenderError(
                    code="output_not_created",
                    message=f"DOCX was not created at {output}",
                    format="docx",
                    stage="save",
                ))

            logger.info(f"DOCX rendered: {output} ({output.stat().st_size} bytes)")
            return Ok(output)

        except Exception as e:
            return Err(RenderError(
                code="render_failed",
                message=f"DOCX rendering failed: {e}",
                format="docx",
                stage="render",
                cause=e,
            ))

    async def _render_basic(
        self,
        content: str,
        output: Path,
        options: dict,
    ) -> Result[Path, DomainError]:
        """Basic HTML to DOCX conversion without htmldocx."""
        try:
            from docx import Document
            from bs4 import BeautifulSoup
        except ImportError as e:
            return Err(RenderError(
                code="dependencies_missing",
                message=f"Required packages not installed: {e}",
                format="docx",
                stage="init",
            ))

        output.parent.mkdir(parents=True, exist_ok=True)

        try:
            doc = Document()
            soup = BeautifulSoup(content, "html.parser")

            # Extract text content with basic formatting
            for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li"]):
                text = element.get_text(strip=True)
                if not text:
                    continue

                if element.name.startswith("h"):
                    level = int(element.name[1])
                    doc.add_heading(text, level=min(level, 9))
                elif element.name == "li":
                    doc.add_paragraph(text, style="List Bullet")
                else:
                    doc.add_paragraph(text)

            doc.save(str(output))

            if not output.exists():
                return Err(RenderError(
                    code="output_not_created",
                    message=f"DOCX was not created at {output}",
                    format="docx",
                    stage="save",
                ))

            logger.info(f"DOCX rendered (basic): {output}")
            return Ok(output)

        except Exception as e:
            return Err(RenderError(
                code="render_failed",
                message=f"DOCX rendering failed: {e}",
                format="docx",
                stage="render",
                cause=e,
            ))
